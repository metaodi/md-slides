"""DOCX converter for md-docs.

Converts a list of document element dicts (produced by the parser's
``parse_markdown_doc``) into a Microsoft Word document using python-docx
and an optional template file.
"""

import importlib.resources
import os

from docx import Document
from docx.enum.text import WD_BREAK

# Mapping from heading level to Word built-in style name.
_HEADING_STYLES = {
    1: "Heading 1",
    2: "Heading 2",
    3: "Heading 3",
    4: "Heading 4",
    5: "Heading 5",
    6: "Heading 6",
}

# Bullet style names (Word built-in).
_BULLET_STYLE = "List Bullet"
_BULLET_STYLE_2 = "List Bullet 2"


def get_default_doc_template_path():
    """Return the absolute path to the bundled default Word template.

    Returns:
        Absolute path string to the packaged template.docx.
    """
    try:
        ref = importlib.resources.files("md_slides").joinpath("template.docx")
        with importlib.resources.as_file(ref) as path:
            return str(path)
    except (AttributeError, TypeError):
        import pkg_resources

        return pkg_resources.resource_filename("md_slides", "template.docx")


def convert_doc(elements, output_path, template_path=None):
    """Convert a list of document element dicts to a Word file.

    Args:
        elements: List of element dicts from ``parse_markdown_doc``.
        output_path: Destination ``.docx`` file path.
        template_path: Path to a ``.docx`` template.  Uses bundled default
            if *None*.

    Returns:
        The *output_path* that was written to.
    """
    if template_path is None:
        template_path = get_default_doc_template_path()

    doc = Document(template_path)

    for elem in elements:
        elem_type = elem.get("type")
        if elem_type == "heading":
            _add_heading(doc, elem)
        elif elem_type == "bullet":
            _add_bullet(doc, elem)
        elif elem_type == "paragraph":
            _add_paragraph(doc, elem)
        elif elem_type == "page_break":
            _add_page_break(doc)

    if not output_path.endswith(".docx"):
        output_path = f"{output_path}.docx"

    doc.save(output_path)
    return output_path


# ── Element helpers ───────────────────────────────────────────────────────────


def _add_heading(doc, elem):
    """Add a heading paragraph with optional inline formatting."""
    level = elem.get("level", 1)
    style_name = _HEADING_STYLES.get(level, "Heading 1")
    para = doc.add_paragraph(style=style_name)
    _apply_runs(para, elem.get("runs", []))


def _add_bullet(doc, elem):
    """Add a bullet-list paragraph."""
    level = elem.get("level", 1)
    style = _BULLET_STYLE if level <= 1 else _BULLET_STYLE_2
    para = doc.add_paragraph(style=style)
    _apply_runs(para, elem.get("runs", []))


def _add_paragraph(doc, elem):
    """Add a normal (body) paragraph."""
    para = doc.add_paragraph()
    _apply_runs(para, elem.get("runs", []))


def _add_page_break(doc):
    """Insert a page break."""
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)


# ── Run formatting ────────────────────────────────────────────────────────────


def _apply_runs(paragraph, runs):
    """Apply formatted runs to a paragraph.

    Args:
        paragraph: python-docx Paragraph object.
        runs: List of run dicts with 'text', 'bold', 'italic'.
    """
    for run_data in runs:
        run = paragraph.add_run(run_data.get("text", ""))
        if run_data.get("bold"):
            run.bold = True
        if run_data.get("italic"):
            run.italic = True
