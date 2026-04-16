"""Tests for md_slides.doc_converter."""

import os
import tempfile

import pytest
from docx import Document
from docx.enum.text import WD_BREAK

from md_slides.doc_converter import convert_doc, get_default_doc_template_path
from md_slides.parser import parse_markdown_doc

# ── Helper ─────────────────────────────────────────────────────────────────────


def _make_docx(md_content, template_path=None):
    """Parse *md_content* and convert it to a temporary DOCX; return the path."""
    elements = parse_markdown_doc(md_content)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as fh:
        out = fh.name
    convert_doc(elements, out, template_path=template_path)
    return out


# ── Template ───────────────────────────────────────────────────────────────────


def test_default_doc_template_exists():
    path = get_default_doc_template_path()
    assert os.path.isfile(path), f"Default Word template not found at {path}"


# ── Headings ───────────────────────────────────────────────────────────────────


def test_heading_level_1():
    out = _make_docx("# My Title")
    try:
        doc = Document(out)
        heading_paras = [p for p in doc.paragraphs if p.style.name == "Heading 1"]
        assert any("My Title" in p.text for p in heading_paras)
    finally:
        os.unlink(out)


def test_heading_level_2():
    out = _make_docx("## Section Title")
    try:
        doc = Document(out)
        heading_paras = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
        assert any("Section Title" in p.text for p in heading_paras)
    finally:
        os.unlink(out)


def test_heading_level_3():
    out = _make_docx("### Sub-section")
    try:
        doc = Document(out)
        heading_paras = [p for p in doc.paragraphs if p.style.name == "Heading 3"]
        assert any("Sub-section" in p.text for p in heading_paras)
    finally:
        os.unlink(out)


def test_multiple_heading_levels():
    md = "# Top\n## Middle\n### Bottom"
    out = _make_docx(md)
    try:
        doc = Document(out)
        h1 = [p for p in doc.paragraphs if p.style.name == "Heading 1"]
        h2 = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
        h3 = [p for p in doc.paragraphs if p.style.name == "Heading 3"]
        assert len(h1) == 1
        assert len(h2) == 1
        assert len(h3) == 1
    finally:
        os.unlink(out)


# ── Bullets ────────────────────────────────────────────────────────────────────


def test_bullet_list():
    out = _make_docx("- First\n- Second\n- Third")
    try:
        doc = Document(out)
        bullet_paras = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
        texts = [p.text for p in bullet_paras]
        assert "First" in texts
        assert "Second" in texts
        assert "Third" in texts
    finally:
        os.unlink(out)


def test_nested_bullet():
    out = _make_docx("- Top level\n    - Nested item")
    try:
        doc = Document(out)
        bullet2_paras = [p for p in doc.paragraphs if p.style.name == "List Bullet 2"]
        assert any("Nested item" in p.text for p in bullet2_paras)
    finally:
        os.unlink(out)


# ── Formatting ─────────────────────────────────────────────────────────────────


def test_bold_text_in_paragraph():
    out = _make_docx("This has **bold** text")
    try:
        doc = Document(out)
        bold_runs = [r for p in doc.paragraphs for r in p.runs if r.bold]
        assert any(r.text == "bold" for r in bold_runs)
    finally:
        os.unlink(out)


def test_italic_text_in_paragraph():
    out = _make_docx("This has *italic* text")
    try:
        doc = Document(out)
        italic_runs = [r for p in doc.paragraphs for r in p.runs if r.italic]
        assert any(r.text == "italic" for r in italic_runs)
    finally:
        os.unlink(out)


def test_bold_text_in_bullet():
    out = _make_docx("- **bold bullet**")
    try:
        doc = Document(out)
        bold_runs = [r for p in doc.paragraphs for r in p.runs if r.bold]
        assert any(r.text == "bold bullet" for r in bold_runs)
    finally:
        os.unlink(out)


def test_italic_text_in_bullet():
    out = _make_docx("- *italic bullet*")
    try:
        doc = Document(out)
        italic_runs = [r for p in doc.paragraphs for r in p.runs if r.italic]
        assert any(r.text == "italic bullet" for r in italic_runs)
    finally:
        os.unlink(out)


def test_bold_text_in_heading():
    out = _make_docx("## Heading with **bold**")
    try:
        doc = Document(out)
        heading_paras = [p for p in doc.paragraphs if p.style.name == "Heading 2"]
        bold_runs = [r for p in heading_paras for r in p.runs if r.bold]
        assert any(r.text == "bold" for r in bold_runs)
    finally:
        os.unlink(out)


# ── Page breaks ────────────────────────────────────────────────────────────────


def test_page_break_inserted():
    md = "# Section 1\n\n---\n\n# Section 2"
    out = _make_docx(md)
    try:
        doc = Document(out)
        # Look for a run containing a page break
        found_break = False
        for para in doc.paragraphs:
            for run in para.runs:
                for br in run.element.findall(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br"
                ):
                    if (
                        br.get(
                            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type"
                        )
                        == "page"
                    ):
                        found_break = True
        assert found_break, "Expected a page break between sections"
    finally:
        os.unlink(out)


# ── Mixed content ──────────────────────────────────────────────────────────────


def test_mixed_content():
    md = "# Document Title\n\nSome intro text.\n\n## Section 1\n\n- Bullet A\n- Bullet B\n\nA paragraph.\n\n---\n\n## Section 2\n\n- Bullet C"
    out = _make_docx(md)
    try:
        doc = Document(out)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Document Title" in all_text
        assert "Some intro text." in all_text
        assert "Section 1" in all_text
        assert "Bullet A" in all_text
        assert "Bullet B" in all_text
        assert "A paragraph." in all_text
        assert "Section 2" in all_text
        assert "Bullet C" in all_text
    finally:
        os.unlink(out)


# ── Template preservation ─────────────────────────────────────────────────────


def test_custom_template_preserves_content(tmp_path):
    """Existing content in a template should be preserved."""
    # Create a template with some pre-existing content
    template = tmp_path / "tpl.docx"
    doc = Document()
    doc.add_paragraph("Existing title page content")
    doc.save(str(template))

    out = str(tmp_path / "out.docx")
    elements = parse_markdown_doc("# New Section\n\n- New bullet")
    convert_doc(elements, out, template_path=str(template))

    result = Document(out)
    all_text = "\n".join(p.text for p in result.paragraphs)
    assert "Existing title page content" in all_text
    assert "New Section" in all_text
    assert "New bullet" in all_text


def test_output_extension_appended():
    """If the output path lacks .docx, the converter appends it."""
    elements = parse_markdown_doc("# Hello")
    with tempfile.NamedTemporaryFile(suffix="", delete=False) as fh:
        base = fh.name
    try:
        result = convert_doc(elements, base)
        assert result.endswith(".docx")
        assert os.path.isfile(result)
    finally:
        if os.path.exists(result):
            os.unlink(result)
        if os.path.exists(base):
            os.unlink(base)
